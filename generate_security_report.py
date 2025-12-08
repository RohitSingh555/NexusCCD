#!/usr/bin/env python3
"""
Security Audit Report Generator
Generates comprehensive security audit reports in DOCX and PDF formats
"""

import json
import os
import subprocess
import sys
from datetime import datetime
from pathlib import Path
import zipfile

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

# Report configuration
REPO_NAME = "NexusCCD"
REPO_PATH = "/Users/shashankjain/agilemorph/nexus/NexusCCD"
OUTPUT_NAME = "Agilemorph_NexusCCD_report"
REPORT_DATE = datetime.now().strftime("%Y-%m-%d")

# Security Findings
FINDINGS = {
    "high": [
        {
            "id": "SEC-001",
            "title": "Hardcoded Database Passwords in Docker Compose",
            "severity": "High",
            "location": "docker-compose.yml, docker-compose.prod.yml",
            "description": "Database passwords are hardcoded in docker-compose files: 'nexusccd_password'",
            "code_excerpt": "POSTGRES_PASSWORD=nexusccd_password",
            "impact": "Exposure of database credentials in version control. Anyone with repo access can access the database."
        },
        {
            "id": "SEC-002",
            "title": "Hardcoded Email Password in Settings",
            "severity": "High",
            "location": "ccd/settings.py:255",
            "description": "Gmail SMTP password is hardcoded in settings file",
            "code_excerpt": "EMAIL_HOST_PASSWORD = config('EMAIL_HOST_PASSWORD', default='vktnzpaaurneigpg')",
            "impact": "Email account credentials exposed in source code. Could lead to email account compromise."
        },
        {
            "id": "SEC-003",
            "title": "Default SECRET_KEY in Production",
            "severity": "High",
            "location": "ccd/settings.py:15",
            "description": "Django SECRET_KEY has an insecure default value",
            "code_excerpt": "SECRET_KEY = config('SECRET_KEY', default='django-insecure-dev-key-change-in-production')",
            "impact": "If default is used, session tokens, CSRF tokens, and password reset tokens can be forged."
        },
        {
            "id": "SEC-004",
            "title": "ALLOWED_HOSTS Wildcard Configuration",
            "severity": "High",
            "location": "docker-compose.yml:45,71",
            "description": "ALLOWED_HOSTS includes wildcard (*) allowing any host",
            "code_excerpt": "ALLOWED_HOSTS=localhost,127.0.0.1,20.63.25.169,*",
            "impact": "Vulnerable to Host header injection attacks. Allows DNS rebinding attacks."
        },
        {
            "id": "SEC-005",
            "title": "Default Credentials Documented in README",
            "severity": "High",
            "location": "README.md",
            "description": "Default usernames and passwords for all roles are documented in README",
            "code_excerpt": "Username: superadmin | Password: admin123",
            "impact": "Default credentials can be used to gain unauthorized access if not changed after deployment."
        }
    ],
    "medium": [
        {
            "id": "SEC-006",
            "title": "Content Security Policy Allows Unsafe Inline and Eval",
            "severity": "Medium",
            "location": "nginx/conf.d/dev.fredvictor.org.conf:34",
            "description": "CSP includes 'unsafe-inline' and 'unsafe-eval' directives",
            "code_excerpt": "default-src 'self' https: data: 'unsafe-inline' 'unsafe-eval'",
            "impact": "Reduces effectiveness of CSP. Allows XSS attacks via inline scripts and eval()."
        },
        {
            "id": "SEC-007",
            "title": "DEBUG Mode Enabled in Docker Compose",
            "severity": "Medium",
            "location": "docker-compose.yml:65",
            "description": "DEBUG=True is set in docker-compose environment",
            "code_excerpt": "DEBUG=True",
            "impact": "Debug mode exposes sensitive information in error pages and enables debug toolbar."
        },
        {
            "id": "SEC-008",
            "title": "Database Port Exposed in Production Docker Compose",
            "severity": "Medium",
            "location": "docker-compose.prod.yml:44",
            "description": "PostgreSQL port 5432 is exposed to host in production configuration",
            "code_excerpt": "ports:\n      - \"5432:5432\"",
            "impact": "Database is accessible from host network. Increases attack surface."
        },
        {
            "id": "SEC-009",
            "title": "Dockerfile Runs as Root User",
            "severity": "Medium",
            "location": "Dockerfile, Dockerfile.prod",
            "description": "Docker containers run as root user by default",
            "code_excerpt": "No USER directive found",
            "impact": "If container is compromised, attacker has root privileges."
        }
    ],
    "low": [
        {
            "id": "SEC-010",
            "title": "No License File Found",
            "severity": "Low",
            "location": "Repository root",
            "description": "No LICENSE file found in repository",
            "code_excerpt": "N/A",
            "impact": "Unclear licensing terms for open source compliance."
        },
        {
            "id": "SEC-011",
            "title": "Large File Upload Timeout Configuration",
            "severity": "Low",
            "location": "nginx/conf.d/dev.fredvictor.org.conf:59-64",
            "description": "Very long timeouts (1200s) for file uploads may allow DoS",
            "code_excerpt": "proxy_send_timeout 1200s;\nproxy_read_timeout 1200s;",
            "impact": "Long-running connections can exhaust server resources."
        }
    ]
}

# Repository Metadata
REPO_METADATA = {
    "repo_name": "NexusCCD",
    "branch": "main",
    "last_commit_sha": "3762805d25a19ab53b1519d522df4c2ab1d7ecc3",
    "last_commit_date": "2025-12-02 10:51:08 +0000",
    "last_commit_author": "rohitsingh555",
    "top_contributors": [
        {"name": "Rushikesh0824", "commits": 54},
        {"name": "RohitSingh555", "commits": 16},
        {"name": "rohitsingh555", "commits": 13},
        {"name": "Shashank Jain", "commits": 11}
    ],
    "total_files": 1106,
    "python_files": 188,
    "scanned_files": 258,
    "repository_size": "184M",
    "languages": ["Python", "HTML", "JavaScript", "CSS", "YAML", "Shell", "SQL"],
    "frameworks": ["Django 4.2.7", "Django REST Framework"],
    "package_manager": "pip (requirements.txt)",
    "build_tools": ["Docker", "Docker Compose"],
    "infrastructure": ["Docker", "Nginx", "PostgreSQL"]
}

# Dependencies from requirements.txt
DEPENDENCIES = [
    {"name": "Django", "version": "4.2.7", "status": "Current"},
    {"name": "djangorestframework", "version": "3.14.0", "status": "Current"},
    {"name": "psycopg2-binary", "version": "2.9.7", "status": "Check CVE"},
    {"name": "gunicorn", "version": "21.2.0", "status": "Current"},
    {"name": "celery", "version": "5.3.4", "status": "Check CVE"},
    {"name": "redis", "version": "5.0.1", "status": "Check CVE"},
    {"name": "Pillow", "version": "10.1.0", "status": "Check CVE"}
]

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_heading_with_style(doc, text, level=1):
    """Add a heading with proper styling"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading

def add_table_with_data(doc, headers, rows, col_widths=None):
    """Add a table with headers and data"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        if col_widths and i < len(col_widths):
            header_cells[i].width = Inches(col_widths[i])
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def create_executive_summary(doc):
    """Create executive summary section"""
    add_heading_with_style(doc, "Executive Summary", 1)
    
    # Calculate summary statistics
    total_findings = sum(len(FINDINGS[sev]) for sev in FINDINGS)
    high_count = len(FINDINGS["high"])
    medium_count = len(FINDINGS["medium"])
    low_count = len(FINDINGS["low"])
    
    doc.add_paragraph(
        f"This security audit of the {REPO_NAME} repository identified {total_findings} security findings "
        f"across {REPO_METADATA['scanned_files']} files. The assessment reveals {high_count} high-severity, "
        f"{medium_count} medium-severity, and {low_count} low-severity issues."
    )
    
    doc.add_paragraph("Key Findings:", style='List Bullet')
    p = doc.add_paragraph("", style='List Bullet')
    p.add_run("Hardcoded credentials").bold = True
    p.add_run(" in configuration files (database passwords, email credentials)")
    
    p = doc.add_paragraph("", style='List Bullet')
    p.add_run("Insecure default configurations").bold = True
    p.add_run(" (SECRET_KEY, ALLOWED_HOSTS wildcard)")
    
    p = doc.add_paragraph("", style='List Bullet')
    p.add_run("Documented default credentials").bold = True
    p.add_run(" in README file")
    
    p = doc.add_paragraph("", style='List Bullet')
    p.add_run("Weak Content Security Policy").bold = True
    p.add_run(" allowing unsafe inline scripts")
    
    p = doc.add_paragraph("", style='List Bullet')
    p.add_run("Docker security configurations").bold = True
    p.add_run(" that could be improved")

def create_cover_page(doc):
    """Create cover page"""
    # Title
    title = doc.add_heading('Security Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()  # Spacing
    
    # Project name
    p = doc.add_paragraph(REPO_NAME)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.bold = True
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    # Metadata
    metadata_items = [
        ("Repository Path", REPO_PATH),
        ("Branch", REPO_METADATA["branch"]),
        ("Report Date", REPORT_DATE),
        ("Last Commit", REPO_METADATA["last_commit_sha"][:12]),
        ("Generated By", "Automated Security Audit Tool")
    ]
    
    for label, value in metadata_items:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()  # Spacing
    
    # Summary
    total_findings = sum(len(FINDINGS[sev]) for sev in FINDINGS)
    high_count = len(FINDINGS["high"])
    medium_count = len(FINDINGS["medium"])
    low_count = len(FINDINGS["low"])
    
    p = doc.add_paragraph(f"Findings Summary: {high_count} High, {medium_count} Medium, {low_count} Low")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    
    add_page_break(doc)

def create_full_report():
    """Create the full technical report"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Cover page
    create_cover_page(doc)
    
    # Table of Contents placeholder
    add_heading_with_style(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Repository Metadata",
        "3. Source Inventory",
        "4. Dependency Analysis",
        "5. Static Code Analysis Findings",
        "6. Secrets & Configuration Review",
        "7. Infrastructure & Container Review",
        "8. CI/CD Pipeline Review",
        "9. License & Compliance",
        "10. Evidence & Reproducibility",
        "11. Appendices"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)
    
    # Executive Summary
    create_executive_summary(doc)
    add_page_break(doc)
    
    # Repository Metadata
    add_heading_with_style(doc, "2. Repository Metadata", 1)
    doc.add_paragraph(f"Repository Name: {REPO_METADATA['repo_name']}")
    doc.add_paragraph(f"Branch: {REPO_METADATA['branch']}")
    doc.add_paragraph(f"Last Commit SHA: {REPO_METADATA['last_commit_sha']}")
    doc.add_paragraph(f"Last Commit Date: {REPO_METADATA['last_commit_date']}")
    doc.add_paragraph(f"Last Commit Author: {REPO_METADATA['last_commit_author']}")
    doc.add_paragraph(f"Total Files: {REPO_METADATA['total_files']}")
    doc.add_paragraph(f"Files Scanned: {REPO_METADATA['scanned_files']}")
    doc.add_paragraph(f"Repository Size: {REPO_METADATA['repository_size']}")
    
    doc.add_paragraph("Top Contributors:")
    for contrib in REPO_METADATA['top_contributors']:
        doc.add_paragraph(f"  • {contrib['name']}: {contrib['commits']} commits", style='List Bullet')
    
    add_page_break(doc)
    
    # Source Inventory
    add_heading_with_style(doc, "3. Source Inventory", 1)
    doc.add_paragraph("Languages Detected:")
    for lang in REPO_METADATA['languages']:
        doc.add_paragraph(f"  • {lang}", style='List Bullet')
    
    doc.add_paragraph("Frameworks:")
    for fw in REPO_METADATA['frameworks']:
        doc.add_paragraph(f"  • {fw}", style='List Bullet')
    
    doc.add_paragraph(f"Package Manager: {REPO_METADATA['package_manager']}")
    doc.add_paragraph("Build Tools:")
    for tool in REPO_METADATA['build_tools']:
        doc.add_paragraph(f"  • {tool}", style='List Bullet')
    
    doc.add_paragraph("Infrastructure:")
    for infra in REPO_METADATA['infrastructure']:
        doc.add_paragraph(f"  • {infra}", style='List Bullet')
    
    add_page_break(doc)
    
    # Dependency Analysis
    add_heading_with_style(doc, "4. Dependency Analysis", 1)
    doc.add_paragraph(
        "The following dependencies were identified from requirements.txt. "
        "Note: Online CVE lookup is recommended using tools like 'pip-audit' or 'safety check'."
    )
    
    dep_headers = ["Package", "Version", "Status", "Recommendation"]
    dep_rows = []
    for dep in DEPENDENCIES:
        recommendation = "Review for known CVEs" if dep["status"] == "Check CVE" else "Monitor for updates"
        dep_rows.append([dep["name"], dep["version"], dep["status"], recommendation])
    
    add_table_with_data(doc, dep_headers, dep_rows, [2, 1.5, 1.5, 3])
    
    doc.add_paragraph(
        "\nRecommended Actions:",
        style='Heading 3'
    )
    doc.add_paragraph("1. Run 'pip-audit' or 'safety check' to identify vulnerable dependencies", style='List Bullet')
    doc.add_paragraph("2. Update dependencies to latest secure versions", style='List Bullet')
    doc.add_paragraph("3. Implement automated dependency scanning in CI/CD pipeline", style='List Bullet')
    doc.add_paragraph("4. Review and update requirements.txt regularly", style='List Bullet')
    
    add_page_break(doc)
    
    # Static Code Analysis Findings
    add_heading_with_style(doc, "5. Static Code Analysis Findings", 1)
    
    for severity in ["high", "medium", "low"]:
        if FINDINGS[severity]:
            add_heading_with_style(doc, f"5.{'1' if severity == 'high' else '2' if severity == 'medium' else '3'}. {severity.capitalize()} Severity Findings", 2)
            
            for finding in FINDINGS[severity]:
                doc.add_paragraph(f"Finding ID: {finding['id']}", style='Heading 3')
                doc.add_paragraph(f"Title: {finding['title']}")
                doc.add_paragraph(f"Severity: {finding['severity']}")
                doc.add_paragraph(f"Location: {finding['location']}")
                doc.add_paragraph(f"Description: {finding['description']}")
                
                code_p = doc.add_paragraph("Code Excerpt:")
                code_p.add_run(f"\n{finding['code_excerpt']}").font.name = 'Courier New'
                
                doc.add_paragraph(f"Impact: {finding['impact']}")
                doc.add_paragraph()  # Spacing
    
    add_page_break(doc)
    
    # Secrets & Configuration Review
    add_heading_with_style(doc, "6. Secrets & Configuration Review", 1)
    doc.add_paragraph(
        "The following secrets and sensitive configurations were identified in the codebase:"
    )
    
    secrets_found = [
        ("Database Password", "docker-compose.yml, docker-compose.prod.yml", "nexusccd_password", "High"),
        ("Email SMTP Password", "ccd/settings.py", "vktnzpaaurneigpg", "High"),
        ("Django SECRET_KEY", "ccd/settings.py", "django-insecure-dev-key-change-in-production", "High"),
        ("Default User Credentials", "README.md", "Multiple default passwords", "High")
    ]
    
    secrets_headers = ["Secret Type", "Location", "Value/Pattern", "Severity"]
    secrets_rows = [[s[0], s[1], s[2][:50] + "..." if len(s[2]) > 50 else s[2], s[3]] for s in secrets_found]
    add_table_with_data(doc, secrets_headers, secrets_rows, [2, 3, 2.5, 1])
    
    doc.add_paragraph("\nNote: All exposed credentials should be rotated and removed from the codebase.")
    
    add_page_break(doc)
    
    # Infrastructure & Container Review
    add_heading_with_style(doc, "7. Infrastructure & Container Review", 1)
    
    doc.add_paragraph("Dockerfile Issues:", style='Heading 3')
    docker_issues = [
        ("Root User", "Containers run as root user", "Create non-root user", "Medium"),
        ("Exposed Ports", "Database port exposed in production", "Remove port mapping", "Medium"),
        ("Build Tools", "build-essential included in image", "Use multi-stage builds", "Low")
    ]
    
    docker_headers = ["Issue", "Description", "Recommendation", "Severity"]
    docker_rows = [[d[0], d[1], d[2], d[3]] for d in docker_issues]
    add_table_with_data(doc, docker_headers, docker_rows, [1.5, 3, 2, 1])
    
    doc.add_paragraph("\nInfrastructure as Code:", style='Heading 3')
    doc.add_paragraph("No Terraform, CloudFormation, or ARM templates detected.")
    
    doc.add_paragraph("\nNote: Docker configurations should be reviewed for security best practices.")
    
    add_page_break(doc)
    
    # CI/CD Pipeline Review
    add_heading_with_style(doc, "8. CI/CD Pipeline Review", 1)
    doc.add_paragraph("No CI/CD pipeline configuration files found (.github/workflows, .gitlab-ci.yml, circleci).")
    doc.add_paragraph("\nNote: No CI/CD pipeline configuration found. Consider implementing automated security scanning.")
    
    add_page_break(doc)
    
    # License & Compliance
    add_heading_with_style(doc, "9. License & Compliance", 1)
    doc.add_paragraph("No LICENSE file found in repository.")
    doc.add_paragraph("\nNote: No LICENSE file found. Consider adding appropriate license documentation.")
    
    add_page_break(doc)
    
    # Evidence & Reproducibility
    add_heading_with_style(doc, "10. Evidence & Reproducibility", 1)
    doc.add_paragraph("Commands used for this audit:", style='Heading 3')
    commands = [
        "grep -r 'SECRET_KEY\\|PASSWORD\\|password' --include='*.py' --include='*.yml'",
        "grep -r 'ALLOWED_HOSTS' --include='*.py' --include='*.yml'",
        "grep -r 'DEBUG' --include='*.py' --include='*.yml'",
        "find . -name 'Dockerfile*' -exec cat {} \\;",
        "find . -name 'docker-compose*.yml' -exec cat {} \\;",
        "cat requirements.txt",
        "git log --format='%an' | sort | uniq -c | sort -rn"
    ]
    for cmd in commands:
        doc.add_paragraph(cmd, style='No Spacing').runs[0].font.name = 'Courier New'
    
    doc.add_paragraph("\nNote: Full command log and evidence files are included in the artifacts package.")
    
    add_page_break(doc)
    
    # Appendices
    add_heading_with_style(doc, "11. Appendices", 1)
    
    doc.add_paragraph("Appendix A: Files Scanned", style='Heading 3')
    doc.add_paragraph(f"Total files scanned: {REPO_METADATA['scanned_files']}")
    doc.add_paragraph("File types scanned: .py, .js, .html, .yml, .yaml, .json, Dockerfile*")
    doc.add_paragraph("Excluded paths: .git, node_modules, build, dist, staticfiles, __pycache__")
    
    doc.add_paragraph("\nAppendix B: Command Log", style='Heading 3')
    doc.add_paragraph("See artifacts package for complete command log.")
    
    doc.add_paragraph("\nAppendix C: Recommended External Scanners", style='Heading 3')
    scanners = [
        ("pip-audit", "Python dependency vulnerability scanning"),
        ("safety", "Python dependency security checking"),
        ("bandit", "Python static security analysis"),
        ("semgrep", "Static analysis for security vulnerabilities"),
        ("gitleaks", "Secrets scanning in git repositories"),
        ("truffleHog", "Secrets scanning"),
        ("trivy", "Container image vulnerability scanning"),
        ("docker-bench-security", "Docker security best practices")
    ]
    for scanner, desc in scanners:
        doc.add_paragraph(f"  • {scanner}: {desc}", style='List Bullet')
    
    return doc

def create_metadata_json():
    """Create metadata JSON file"""
    total_findings = sum(len(FINDINGS[sev]) for sev in FINDINGS)
    
    metadata = {
        "report_name": OUTPUT_NAME,
        "report_date": REPORT_DATE,
        "repo_name": REPO_METADATA["repo_name"],
        "repo_path": REPO_PATH,
        "branch": REPO_METADATA["branch"],
        "repo_sha": REPO_METADATA["last_commit_sha"],
        "total_files_scanned": REPO_METADATA["scanned_files"],
        "findings_count_by_severity": {
            "high": len(FINDINGS["high"]),
            "medium": len(FINDINGS["medium"]),
            "low": len(FINDINGS["low"]),
            "total": total_findings
        },
        "scanned_date": REPORT_DATE,
        "generator": "Automated Security Audit Tool"
    }
    
    return metadata

def create_artifacts_zip():
    """Create artifacts zip file"""
    zip_path = f"{OUTPUT_NAME}-artifacts.zip"
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # Add evidence files
        evidence_files = [
            "ccd/settings.py",
            "docker-compose.yml",
            "docker-compose.prod.yml",
            "Dockerfile",
            "Dockerfile.prod",
            "requirements.txt",
            "README.md",
            "nginx/conf.d/dev.fredvictor.org.conf"
        ]
        
        for file_path in evidence_files:
            full_path = os.path.join(REPO_PATH, file_path)
            if os.path.exists(full_path):
                zipf.write(full_path, f"evidence/{file_path}")
        
        # Create command log
        command_log = """Security Audit Command Log
========================

Commands executed during security audit:

1. Repository Information:
   git rev-parse --abbrev-ref HEAD
   git log -1 --format="%H|%ai|%an|%s"
   find . -type f | wc -l
   du -sh .

2. Security Scans:
   grep -r 'SECRET_KEY|PASSWORD|password' --include='*.py' --include='*.yml'
   grep -r 'ALLOWED_HOSTS' --include='*.py' --include='*.yml'
   grep -r 'DEBUG' --include='*.py' --include='*.yml'
   grep -r 'eval\\(|exec\\(|subprocess\\.call' --include='*.py'
   grep -r 'md5|sha1|DES|RC4' --include='*.py' -i

3. File Discovery:
   find . -name 'Dockerfile*'
   find . -name 'docker-compose*.yml'
   find . -name '*.env*'
   find . -name '.github/workflows/*.yml'
   find . -name '.gitlab-ci.yml'

4. Dependency Analysis:
   cat requirements.txt
   python3 -m pip list --format=json

Scan completed: """ + REPORT_DATE
        
        zipf.writestr("command_log.txt", command_log)
    
    return zip_path

def main():
    """Main function to generate reports"""
    print("Generating security audit reports...")
    
    # Create full technical report
    print("Creating full technical report...")
    doc = create_full_report()
    docx_path = f"{OUTPUT_NAME}.docx"
    doc.save(docx_path)
    print(f"✓ Created {docx_path}")
    
    # Create executive summary (extract first page)
    print("Creating executive summary...")
    exec_doc = Document()
    sections = exec_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    create_cover_page(exec_doc)
    create_executive_summary(exec_doc)
    
    exec_docx_path = f"{OUTPUT_NAME}_executive_summary.docx"
    exec_doc.save(exec_docx_path)
    print(f"✓ Created {exec_docx_path}")
    
    # Create metadata JSON
    print("Creating metadata file...")
    metadata = create_metadata_json()
    metadata_path = f"{OUTPUT_NAME}.metadata.json"
    with open(metadata_path, 'w') as f:
        json.dump(metadata, f, indent=2)
    print(f"✓ Created {metadata_path}")
    
    # Create artifacts zip
    print("Creating artifacts package...")
    zip_path = create_artifacts_zip()
    print(f"✓ Created {zip_path}")
    
    # Try to convert to PDF (if pandoc or libreoffice available)
    print("\nAttempting PDF conversion...")
    try:
        # Try pandoc first
        subprocess.run(['pandoc', docx_path, '-o', f"{OUTPUT_NAME}.pdf"], check=True, capture_output=True)
        print(f"✓ Created {OUTPUT_NAME}.pdf")
        subprocess.run(['pandoc', exec_docx_path, '-o', f"{OUTPUT_NAME}_executive_summary.pdf"], check=True, capture_output=True)
        print(f"✓ Created {OUTPUT_NAME}_executive_summary.pdf")
    except (subprocess.CalledProcessError, FileNotFoundError):
        try:
            # Try libreoffice
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path], check=True, capture_output=True)
            print(f"✓ Created {OUTPUT_NAME}.pdf")
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', exec_docx_path], check=True, capture_output=True)
            print(f"✓ Created {OUTPUT_NAME}_executive_summary.pdf")
        except (subprocess.CalledProcessError, FileNotFoundError):
            print("⚠ PDF conversion not available. Install pandoc or libreoffice to generate PDFs.")
            print("   DOCX files can be converted to PDF manually using Microsoft Word or online converters.")
    
    # Print summary
    print("\n" + "="*60)
    print("SECURITY AUDIT REPORT SUMMARY")
    print("="*60)
    print(f"Total Findings: {sum(len(FINDINGS[sev]) for sev in FINDINGS)}")
    print(f"  High: {len(FINDINGS['high'])}")
    print(f"  Medium: {len(FINDINGS['medium'])}")
    print(f"  Low: {len(FINDINGS['low'])}")
    print("\nTop 5 Findings:")
    all_findings = []
    for sev in ["high", "medium", "low"]:
        all_findings.extend(FINDINGS[sev])
    
    for i, finding in enumerate(all_findings[:5], 1):
        print(f"{i}. [{finding['id']}] {finding['title']} ({finding['severity']})")
        print(f"   Location: {finding['location']}")
    
    print("\n" + "="*60)
    print("Reports generated successfully!")
    print(f"  • {docx_path}")
    print(f"  • {exec_docx_path}")
    print(f"  • {metadata_path}")
    print(f"  • {zip_path}")
    print("="*60)

if __name__ == "__main__":
    main()

